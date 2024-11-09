# -*- coding: utf-8 -*-
"""
ICTPRG443 Assessment 3

Created on Mon Oct 28 18:08:13 2024

@author: Bailey
"""

from docx import Document
from docx.shared import Inches
import os

# Creates an empty word doc
def create_word_document(name):
    doc = Document()
    doc.save(name)

# Adds a heading to a word doc
def add_heading(document, heading, level):
    doc = Document(document)
    doc.add_heading(heading, level)
    doc.save(document)

# Adds a picture to a word doc
def add_picture(document, picture, width, height):
    doc = Document(document)
    doc.add_picture(picture, width=width, height=height)
    doc.save(document)

# Adds a paragraph to a word doc
def add_paragraph(document, paragraph, style):
    doc = Document(document)
    doc.add_paragraph(paragraph, style = style)
    doc.save(document)

# returns the text in a word doc
def get_text_from_doc(document):
    doc = Document(document)
    text = []
    for line in doc.paragraphs:
        text.append(line.text)
    return '\n'.join(text)



def main():
    print('Word Manager Program by Bailey Byl')
    document_path = None
    while True:
        print('''Options:
        S - Select document
        C - Create document
        H - Add a heading
        P - Add a paragraph
        I - Add an image
        G - Get all text in the document
        ENTER to exit''')

        option = input('Select option: ')

        # Select a word doc
        if option == 'S' or option == 's':
            selected_document = input('Enter document name: ')
            path = './' + selected_document
            if os.path.exists(path):
                document_path = selected_document
                print('Document selected')
            else:
                print('Document not found')

        # Create a word doc
        elif option == 'C' or option == 'c':
            document_name = input('Enter document name: ')
            try:
                create_word_document(document_name)
                document_path = document_name
                print('Document created and selected')
            except:
                print('Document failed to create, may already exit or incorrect name')

        # Add a heading to a word doc
        elif option == 'H' or option == 'h':
            if document_path is None:
                print('No document selected, select one first')
            else:
                heading = input('Enter heading: ')
                level = int(input('Enter level (1,2 etc): '))
                try:
                    add_heading(document_path, heading, level)
                    print('Heading added to document')
                except:
                    print('Heading not added, check if the level is valid')

        # Add a paragraph to a word doc
        elif option == 'P' or option == 'p':
            if document_path is None:
                print('No document selected, select one first')
            else:
                paragraph = input('Enter paragraph: ')
                style = input('Enter style (Press ENTER for nothing): ')
                if style == '':
                    style = None
                try:
                    add_paragraph(document_path, paragraph, style)
                    print('paragraph added to document')
                except:
                    print('paragraph not added, check if the style is valid')

        # Add an image to a word doc
        elif option == 'I' or option == 'i':
            if document_path is None:
                print('No document selected, select one first')
            else:
                try:
                    image = input('Enter image name: ')
                    width = int(input('Enter width: '))
                    height = int(input('Enter height: '))
                    add_picture(document_path, image, Inches(width), Inches(height))
                    print('Image added to document')
                except:
                    print('Image not added, check if the width, height or image name is valid')

        # Prints the text in a word doc
        elif option == 'G' or option == 'g':
            if document_path is None:
                print('No document selected, select one first')
            else:
                print(get_text_from_doc(document_path))

        # Exits the program
        elif option == '':
            return

        else:
            print('Invalid input')



if __name__ == '__main__':
    main()



