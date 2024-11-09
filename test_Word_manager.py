# -*- coding: utf-8 -*-
"""
ICTPRG443 Assessment 3

Created on Mon Oct 28 18:08:36 2024

@author: Bailey Byl 20135558
"""

import unittest
import os
import docx

from Word_manager import add_heading, add_paragraph, add_picture, create_word_document, get_text_from_doc


class TestPdfManagerFunctions(unittest.TestCase):

    def test_create_word_document(self):
        create_word_document('test_create.docx')
        self.assertTrue(os.path.isfile('test_create.docx'))

    def test_add_heading(self):
        create_word_document('test_heading.docx')
        add_heading('test_heading.docx', 'Test Heading', 1)
        self.assertEqual('Test Heading', get_text_from_doc('test_heading.docx'))

    def test_add_paragraph(self):
        create_word_document('test_paragraph.docx')
        add_paragraph('test_paragraph.docx', 'Test paragraph', None)
        self.assertEqual('Test paragraph', get_text_from_doc('test_paragraph.docx'))

    def test_add_picture(self):
        create_word_document('test_picture.docx')
        try:
            add_picture('test_picture.docx', 'image.png', 1, 1)
        except:
            self.fail('Failed to add picture')




if __name__ == '__main__':
    unittest.main()
